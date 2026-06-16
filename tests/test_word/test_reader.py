"""Unit tests for the Word document reader."""

import io
from pathlib import Path

import pytest
from docx import Document

from backend.app.word.models import ElementType, WordDocument
from backend.app.word.reader import WordReadError, read_docx


def _make_docx(paragraphs: list[tuple[str, str]] | None = None) -> bytes:
    """
    Build a minimal .docx in memory.
    paragraphs: list of (text, style_name) tuples. Style must be a built-in Word style.
    """
    doc = Document()
    if paragraphs:
        for text, style in paragraphs:
            doc.add_paragraph(text, style=style)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Basic parsing ─────────────────────────────────────────────────────────────

def test_read_docx_from_bytes() -> None:
    raw = _make_docx([("Hello world", "Normal")])
    doc = read_docx(raw)
    assert isinstance(doc, WordDocument)
    assert doc.source_filename == "<bytes>"


def test_read_docx_from_path(tmp_path: Path) -> None:
    raw = _make_docx([("Hello world", "Normal")])
    p = tmp_path / "test.docx"
    p.write_bytes(raw)
    doc = read_docx(p)
    assert doc.source_filename == "test.docx"


def test_missing_file_raises(tmp_path: Path) -> None:
    with pytest.raises(WordReadError):
        read_docx(tmp_path / "nonexistent.docx")


def test_invalid_bytes_raises() -> None:
    with pytest.raises(WordReadError):
        read_docx(b"not a docx file at all")


# ── Sections ─────────────────────────────────────────────────────────────────

def test_heading1_creates_section() -> None:
    raw = _make_docx([
        ("Introduction", "Heading 1"),
        ("Some body text.", "Normal"),
    ])
    doc = read_docx(raw)
    assert len(doc.sections) == 1
    assert doc.sections[0].title == "Introduction"


def test_multiple_sections() -> None:
    raw = _make_docx([
        ("Section A", "Heading 1"),
        ("Body A.", "Normal"),
        ("Section B", "Heading 1"),
        ("Body B.", "Normal"),
    ])
    doc = read_docx(raw)
    assert len(doc.sections) == 2
    assert doc.sections[0].title == "Section A"
    assert doc.sections[1].title == "Section B"


def test_preamble_before_first_heading() -> None:
    raw = _make_docx([
        ("Preamble text before any heading.", "Normal"),
        ("Section One", "Heading 1"),
    ])
    doc = read_docx(raw)
    assert any(el.text == "Preamble text before any heading." for el in doc.preamble)


# ── Element types ─────────────────────────────────────────────────────────────

def test_paragraph_element_type() -> None:
    raw = _make_docx([
        ("My Section", "Heading 1"),
        ("A normal paragraph.", "Normal"),
    ])
    doc = read_docx(raw)
    paras = [e for e in doc.sections[0].elements if e.element_type == ElementType.PARAGRAPH]
    assert len(paras) == 1
    assert paras[0].text == "A normal paragraph."


def test_subheading_element() -> None:
    raw = _make_docx([
        ("Main", "Heading 1"),
        ("Sub", "Heading 2"),
    ])
    doc = read_docx(raw)
    subs = [e for e in doc.sections[0].elements if e.element_type == ElementType.HEADING_2]
    assert len(subs) == 1
    assert subs[0].text == "Sub"


def test_table_element() -> None:
    docx_doc = Document()
    docx_doc.add_paragraph("Section", "Heading 1")
    tbl = docx_doc.add_table(rows=2, cols=3)
    tbl.cell(0, 0).text = "H1"
    tbl.cell(0, 1).text = "H2"
    tbl.cell(0, 2).text = "H3"
    tbl.cell(1, 0).text = "A"
    tbl.cell(1, 1).text = "B"
    tbl.cell(1, 2).text = "C"
    buf = io.BytesIO()
    docx_doc.save(buf)

    doc = read_docx(buf.getvalue())
    tables = [e for e in doc.sections[0].elements if e.element_type == ElementType.TABLE]
    assert len(tables) == 1
    assert tables[0].table is not None
    assert tables[0].table.rows == 2
    assert tables[0].table.cols == 3


# ── Word count ────────────────────────────────────────────────────────────────

def test_total_word_count() -> None:
    raw = _make_docx([
        ("My Section", "Heading 1"),
        ("one two three four five", "Normal"),
    ])
    doc = read_docx(raw)
    # "My Section" (2) + "one two three four five" (5) = 7
    assert doc.total_word_count == 7


def test_empty_document_word_count() -> None:
    raw = _make_docx()
    doc = read_docx(raw)
    assert doc.total_word_count == 0
