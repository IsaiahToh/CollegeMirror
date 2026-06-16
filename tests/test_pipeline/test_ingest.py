"""Integration tests for the ingestion pipeline (Claude mocked)."""

import io
from pathlib import Path
from unittest.mock import MagicMock, patch

from docx import Document

from backend.app.models.design_schema import ContentMapping, ExampleAnalysis, SemanticRole


def _make_docx(tmp_path: Path, name: str = "source.docx") -> Path:
    doc = Document()
    doc.add_paragraph("Executive Summary", "Heading 1")
    doc.add_paragraph("This is the body text of the document.", "Normal")
    buf = io.BytesIO()
    doc.save(buf)
    p = tmp_path / name
    p.write_bytes(buf.getvalue())
    return p


def _mock_analysis(idml_filename: str, word_filenames: list[str]) -> ExampleAnalysis:
    return ExampleAnalysis(
        idml_filename=idml_filename,
        word_filenames=word_filenames,
        mappings=[
            ContentMapping(
                word_element_type="heading_1",
                word_text_preview="Executive Summary",
                idml_frame_id="tf_0001",
                idml_spread_index=0,
                role=SemanticRole.SECTION_HEADER,
                paragraph_style="ParagraphStyle/Body Text",
            )
        ],
        spread_count=1,
        page_count=1,
    )


async def test_run_ingestion_single_pair(tmp_path: Path, minimal_idml_file: Path) -> None:
    """Ingestion with one IDML + one Word doc produces a valid DesignSchema."""
    from backend.app.core.storage import Storage
    from backend.app.pipeline.ingest import run_ingestion

    word_path = _make_docx(tmp_path)
    analysis = _mock_analysis(minimal_idml_file.name, [word_path.name])

    storage = MagicMock(spec=Storage)
    storage.example_set_dir.return_value = tmp_path

    with patch("backend.app.pipeline.ingest.map_content", return_value=analysis):
        schema = await run_ingestion(
            example_set_id="test_set",
            pairs=[(minimal_idml_file, [word_path])],
            storage=storage,
        )

    assert schema.example_set_id == "test_set"
    assert len(schema.example_analyses) == 1
    assert schema.example_analyses[0].word_filenames == [word_path.name]
    assert schema.primary_template_idml == minimal_idml_file.name
    storage.save_schema.assert_called_once()


async def test_run_ingestion_multi_word(tmp_path: Path, minimal_idml_file: Path) -> None:
    """One IDML paired with two Word docs — both filenames appear in the analysis."""
    from backend.app.core.storage import Storage
    from backend.app.pipeline.ingest import run_ingestion

    word_a = _make_docx(tmp_path, "doc_a.docx")
    word_b = _make_docx(tmp_path, "doc_b.docx")
    analysis = _mock_analysis(minimal_idml_file.name, ["doc_a.docx", "doc_b.docx"])

    storage = MagicMock(spec=Storage)
    storage.example_set_dir.return_value = tmp_path

    with patch("backend.app.pipeline.ingest.map_content", return_value=analysis) as mock_map:
        schema = await run_ingestion(
            example_set_id="multi_test",
            pairs=[(minimal_idml_file, [word_a, word_b])],
            storage=storage,
        )

    # map_content should have been called once with a list of 2 docs
    call_args = mock_map.call_args
    word_docs_passed = call_args[0][0]  # first positional arg
    assert len(word_docs_passed) == 2
    assert schema.example_analyses[0].word_filenames == ["doc_a.docx", "doc_b.docx"]


async def test_run_ingestion_multiple_pairs(tmp_path: Path, minimal_idml_file: Path) -> None:
    """Two IDML+Word pairs → two analyses, frame templates derived from both."""
    from backend.app.core.storage import Storage
    from backend.app.pipeline.ingest import run_ingestion

    word_a = _make_docx(tmp_path, "doc_a.docx")
    word_b = _make_docx(tmp_path, "doc_b.docx")

    idml_b = tmp_path / "second.idml"
    idml_b.write_bytes(minimal_idml_file.read_bytes())

    analyses = [
        _mock_analysis(minimal_idml_file.name, ["doc_a.docx"]),
        _mock_analysis("second.idml", ["doc_b.docx"]),
    ]

    storage = MagicMock(spec=Storage)
    storage.example_set_dir.return_value = tmp_path

    call_count = 0

    def side_effect(*args: object, **kwargs: object) -> ExampleAnalysis:
        nonlocal call_count
        result = analyses[call_count]
        call_count += 1
        return result

    with patch("backend.app.pipeline.ingest.map_content", side_effect=side_effect):
        schema = await run_ingestion(
            example_set_id="two_pairs",
            pairs=[
                (minimal_idml_file, [word_a]),
                (idml_b, [word_b]),
            ],
            storage=storage,
        )

    assert len(schema.example_analyses) == 2
    assert schema.primary_template_idml == minimal_idml_file.name
    storage.save_schema.assert_called_once()
